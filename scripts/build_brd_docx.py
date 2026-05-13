from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE_COLOR = RGBColor(31, 64, 104)
ACCENT_COLOR = RGBColor(52, 99, 153)
TEXT_COLOR = RGBColor(40, 40, 40)
LIGHT_FILL = "EAF1FB"
HEADER_FILL = "D9E4F5"
BORDER_COLOR = "9FB6D5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = BORDER_COLOR, size: int = 6) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
      tc_borders = OxmlElement("w:tcBorders")
      tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_page_layout(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.4)
    sec.footer_distance = Inches(0.35)


def set_run_font(run, *, name="Aptos", size=10.5, bold=False, color=TEXT_COLOR, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def ensure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, font_size, color, bold in [
        ("Title", 23, TITLE_COLOR, True),
        ("Subtitle", 12, ACCENT_COLOR, False),
        ("Heading 1", 15, TITLE_COLOR, True),
        ("Heading 2", 12.5, ACCENT_COLOR, True),
        ("Heading 3", 11, TITLE_COLOR, True),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(font_size)
        style.font.color.rgb = color
        style.font.bold = bold

    if "Table Header" not in [s.name for s in doc.styles]:
        th = doc.styles.add_style("Table Header", WD_STYLE_TYPE.PARAGRAPH)
        th.font.name = "Aptos"
        th.font.size = Pt(9.5)
        th.font.bold = True
        th.font.color.rgb = TITLE_COLOR

    if "Small Note" not in [s.name for s in doc.styles]:
        sn = doc.styles.add_style("Small Note", WD_STYLE_TYPE.PARAGRAPH)
        sn.font.name = "Aptos"
        sn.font.size = Pt(9)
        sn.font.color.rgb = RGBColor(90, 90, 90)


def add_header_footer(doc: Document, project_name: str) -> None:
    for sec in doc.sections:
        header = sec.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp_run = hp.add_run(project_name)
        set_run_font(hp_run, size=8.5, color=ACCENT_COLOR)

        footer = sec.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp_run = fp.add_run("Business Requirements Document")
        set_run_font(fp_run, size=8.5, color=ACCENT_COLOR)


def add_cover_page(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    p.add_run(title)

    p2 = doc.add_paragraph(style="Subtitle")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(14)
    p2.add_run(subtitle)

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_table.autofit = False
    rows = [
        ("Project", "Excel Template App"),
        ("Document Type", "Business Requirements Document"),
        ("Prepared Date", "2026-05-13"),
        ("Prepared By", "Codex"),
    ]
    for i, (k, v) in enumerate(rows):
        info_table.rows[i].cells[0].width = Inches(1.8)
        info_table.rows[i].cells[1].width = Inches(4.8)
        for j, text in enumerate((k, v)):
            cell = info_table.rows[i].cells[j]
            cell.text = ""
            pcell = cell.paragraphs[0]
            run = pcell.add_run(text)
            set_run_font(run, bold=(j == 0))
            set_cell_border(cell, size=8)
            if j == 0:
                set_cell_shading(cell, LIGHT_FILL)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph("")
    sec = doc.sections[0]
    sect_pr = sec._sectPr
    type_el = sect_pr.xpath("./w:type")
    if type_el:
        type_el[0].set(qn("w:val"), "nextPage")
    else:
        new_type = OxmlElement("w:type")
        new_type.set(qn("w:val"), "nextPage")
        sect_pr.append(new_type)


def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("*", "")
    return text.strip()


def parse_table(lines: list[str], start: int):
    block = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        block.append(lines[i].rstrip())
        i += 1
    if len(block) < 2:
        return None, start + 1
    rows = []
    for idx, line in enumerate(block):
        if idx == 1 and re.match(r"^\|[\-\| :]+\|?$", line.strip()):
            continue
        parts = [clean_inline(x) for x in line.strip().strip("|").split("|")]
        rows.append(parts)
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            value = row[c_idx] if c_idx < len(row) else ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_run_font(run, size=9.5, bold=(r_idx == 0), color=TITLE_COLOR if r_idx == 0 else TEXT_COLOR)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell, size=8 if r_idx == 0 else 6)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "F8FBFF")
    doc.add_paragraph("")


def add_bullet(doc: Document, text: str, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(clean_inline(text))
    set_run_font(run)


def add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(clean_inline(text))
    set_run_font(run, size=9 if style == "Small Note" else 10.5)


def build_docx(markdown_path: Path, output_path: Path) -> None:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_page_layout(doc)
    ensure_styles(doc)

    title = "Business Requirements Document"
    subtitle = markdown_path.stem.replace("_", " ")
    add_cover_page(doc, title, subtitle)
    add_header_footer(doc, subtitle)

    i = 0
    skip_first_h1 = True
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, next_i = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
                i = next_i
                continue

        if stripped.startswith("# "):
            if skip_first_h1:
                skip_first_h1 = False
            else:
                add_paragraph(doc, stripped[2:], "Heading 1")
            i += 1
            continue
        if stripped.startswith("## "):
            add_paragraph(doc, stripped[3:], "Heading 1")
            i += 1
            continue
        if stripped.startswith("### "):
            add_paragraph(doc, stripped[4:], "Heading 2")
            i += 1
            continue
        if stripped.startswith("#### "):
            add_paragraph(doc, stripped[5:], "Heading 3")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_bullet(doc, re.sub(r"^\d+\.\s+", "", stripped), numbered=True)
            i += 1
            continue

        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:], numbered=False)
            i += 1
            continue

        add_paragraph(doc, stripped)
        i += 1

    doc.save(output_path)


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print("Usage: build_brd_docx.py <input.md> <output.docx>")
        return 1
    input_path = Path(argv[1])
    output_path = Path(argv[2])
    build_docx(input_path, output_path)
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
